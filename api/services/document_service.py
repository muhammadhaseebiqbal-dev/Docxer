from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import docx.oxml
import docx.oxml.ns
import os
import re
from utils.config import settings

class DocumentService:
    def __init__(self):
        self.output_dir = settings.output_dir
    
    def create_document(self, markdown_content: str, task_id: str) -> str:
        """Convert markdown content to Word document with proper formatting"""
        doc = Document()
        self._setup_document_styles(doc)
        
        # Add title with better formatting
        title = doc.add_heading('Code Documentation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a subtitle with file info
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Generated with AI-Powered Analysis')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.italic = True
        
        # Add separator
        doc.add_paragraph().add_run('─' * 50).font.color.rgb = RGBColor(200, 200, 200)
        
        # Process markdown content line by line
        lines = markdown_content.split('\n')
        current_code_block = []
        in_code_block = False
        current_language = ''
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                if not in_code_block:
                    doc.add_paragraph()
                i += 1
                continue
            
            # Handle code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End of code block
                    self._add_formatted_code_block(doc, '\n'.join(current_code_block), current_language)
                    current_code_block = []
                    in_code_block = False
                    current_language = ''
                else:
                    # Start of code block
                    in_code_block = True
                    current_language = line[3:].strip()  # Get language after ```
                i += 1
                continue
            
            if in_code_block:
                current_code_block.append(line)
                i += 1
                continue
            
            # Handle headings with better formatting
            if line.startswith('**') and line.endswith('**'):
                heading_text = line.strip('*').strip()
                heading = doc.add_heading(heading_text, level=1)
                self._format_heading(heading, 1)
            elif line.startswith('### '):
                heading_text = line.replace('### ', '').strip()
                heading = doc.add_heading(heading_text, level=3)
                self._format_heading(heading, 3)
            elif line.startswith('## '):
                heading_text = line.replace('## ', '').strip()
                heading = doc.add_heading(heading_text, level=2)
                self._format_heading(heading, 2)
            elif line.startswith('# '):
                heading_text = line.replace('# ', '').strip()
                heading = doc.add_heading(heading_text, level=1)
                self._format_heading(heading, 1)
            # Handle bullet points
            elif line.startswith('- ') or line.startswith('* '):
                bullet_text = line[2:].strip()
                bullet_text = self._clean_markdown(bullet_text)
                para = doc.add_paragraph(bullet_text, style='List Bullet')
                self._format_bullet_point(para)
            # Handle numbered lists
            elif re.match(r'^\d+\.\s', line):
                list_text = re.sub(r'^\d+\.\s', '', line).strip()
                list_text = self._clean_markdown(list_text)
                para = doc.add_paragraph(list_text, style='List Number')
                self._format_numbered_list(para)
            # Handle regular paragraphs
            else:
                # Clean up markdown formatting
                clean_text = self._clean_markdown(line)
                if clean_text:
                    para = doc.add_paragraph(clean_text)
                    self._format_paragraph(para)
            
            i += 1
        
        # Handle any remaining code block
        if current_code_block:
            self._add_formatted_code_block(doc, '\n'.join(current_code_block), current_language)
        
        # Add footer
        self._add_footer(doc)
        
        # Save document
        filename = f"documentation_{task_id}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def _setup_document_styles(self, doc: Document):
        """Set up custom styles for the document"""
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Create code style if it doesn't exist
        try:
            code_style = doc.styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
            code_font = code_style.font
            code_font.name = 'Consolas'
            code_font.size = Pt(9)
            code_style.paragraph_format.left_indent = Inches(0.25)
            code_style.paragraph_format.right_indent = Inches(0.25)
            code_style.paragraph_format.space_before = Pt(6)
            code_style.paragraph_format.space_after = Pt(6)
        except:
            pass  # Style might already exist
    
    def _add_formatted_code_block(self, doc: Document, code_content: str, language: str = ''):
        """Add a properly formatted code block"""
        # Add language label if provided
        if language:
            lang_para = doc.add_paragraph()
            run = lang_para.add_run(f'Language: {language.upper()}')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.bold = True
            lang_para.paragraph_format.space_before = Pt(12)
            lang_para.paragraph_format.space_after = Pt(3)
        
        # Create a table for the code block with better styling
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Light Grid Accent 1'  # Use a lighter table style
        
        cell = table.cell(0, 0)
        
        # Add code content with proper formatting
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(code_content)
        
        # Format the code text
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black text instead of white
        
        # Set cell background to light gray instead of dark
        shading_elm = docx.oxml.parse_xml(
            r'<w:shd {} w:fill="F8F8F8"/>'.format(docx.oxml.ns.nsdecls('w'))
        )
        cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Add padding to the cell
        cell_margins = cell._element.get_or_add_tcPr()
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin_elem = docx.oxml.parse_xml(
                f'<w:tcMar {docx.oxml.ns.nsdecls("w")}>'
                f'<w:{margin_name} w:w="120" w:type="dxa"/>'
                f'</w:tcMar>'
            )
            cell_margins.append(margin_elem)
        
        # Add some space after the code block
        doc.add_paragraph()
    
    def _format_heading(self, heading, level):
        """Format headings with consistent styling"""
        if level == 1:
            heading.runs[0].font.color.rgb = RGBColor(31, 78, 121)  # Dark blue
            heading.runs[0].font.size = Pt(16)
        elif level == 2:
            heading.runs[0].font.color.rgb = RGBColor(79, 129, 189)  # Medium blue
            heading.runs[0].font.size = Pt(14)
        elif level == 3:
            heading.runs[0].font.color.rgb = RGBColor(128, 128, 128)  # Gray
            heading.runs[0].font.size = Pt(12)
        
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    
    def _format_paragraph(self, paragraph):
        """Format regular paragraphs"""
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.15
        
        # Handle inline code formatting
        for run in paragraph.runs:
            if '`' in run.text:
                # Simple inline code formatting
                run.text = run.text.replace('`', '')
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(199, 37, 78)  # Red color for inline code
    
    def _format_bullet_point(self, paragraph):
        """Format bullet points"""
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.left_indent = Inches(0.25)
    
    def _format_numbered_list(self, paragraph):
        """Format numbered lists"""
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.left_indent = Inches(0.25)
    
    def _clean_markdown(self, text: str) -> str:
        """Clean markdown formatting from text"""
        # Remove bold formatting
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        # Remove italic formatting  
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        # Handle inline code - don't remove backticks here, handle in _format_paragraph
        # Remove extra whitespace
        text = ' '.join(text.split())
        
        return text
    
    def _add_footer(self, doc: Document):
        """Add a professional footer"""
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = footer_para.add_run('Generated by Docxer - AI-Powered Code Documentation')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.italic = True
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        from datetime import datetime
        date_run = date_para.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        date_run.font.size = Pt(8)
        date_run.font.color.rgb = RGBColor(150, 150, 150)
