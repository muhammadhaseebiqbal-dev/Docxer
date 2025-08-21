from tkinter import Tk, Label, Button, StringVar
from tkinter.filedialog import askopenfilename
import google.generativeai as genai
import docx
from docx import Document

# This is the GUI code for the docxer app
app = Tk()
app.title("Docxer")
app.geometry("250x250")
app.resizable(False, False)

# Input State
inputState = StringVar()
inputState.set("...........")

# Program's Processing logs
logs = StringVar()
logs.set(".........................................................")

# Field Colors
fieldColor=("#ff8080")

filePath = ""
rawData = ""

genai.configure(api_key="place your api key here")

# Function to read code file
def readCodefile(path):
    global rawData
    with open(path,'r') as f:
        rawData = f.read()
        return True

# Function to open file
def selectFile():
    global filePath
    tempPath = askopenfilename()
    filePath = tempPath
    if tempPath:
        inputState.set("Seleted!")
        print(tempPath)
        fieldColor = "#80FF80"
        fileName.configure(bg=fieldColor, highlightbackground=fieldColor, highlightcolor=fieldColor, highlightthickness=1)
        import threading
        threading.Thread(target=generateDocx).start()

def generateDocx():
    if filePath:
        logs.set("Processing...")
        if readCodefile(filePath):
            logs.set("Document readed successfully!")
        else:
            logs.set("Error reading file!")
        if rawData:
            if askLLM(rawData):
                logs.set("Documentation created!")
            else:
                logs.set("LLM Server Error!")
        else:
            logs.set("Empty/Courupted File!")
                
    else:
        logs.set("Please select a file first!")

def askLLM(rawinput):

    generation_config = {
        "temperature": 1,
        "top_p": 0.95,
        "top_k": 40,
        "max_output_tokens": 8192,
        "response_mime_type": "text/plain",
    }

    model = genai.GenerativeModel(
    model_name="gemini-2.0-flash-exp",
    generation_config=generation_config,
    )

    chat_session = model.start_chat(
    history=[
        {
            "role": "user",
            "parts": [
                "You, are the LLM that create documentation of the provided code, right?\n",
        ],
        },
        {
            "role": "model",
            "parts": [
                "Yes, I am the LLM that creates documentation of the provided code.\n",
        ],
        },
    ]
    )

    response = chat_session.send_message(rawinput)

    parsed_response = response.text.split("\n");


    doc = Document()
    for i in parsed_response:
        if i.__contains__("**"):
            doc.add_heading((((i.removeprefix("**").removesuffix("**")).removeprefix("* ")).replace("**"," ")).removeprefix("`").removesuffix('`'), level=1)
        elif i.__contains__("```"):
            table = doc.add_table(1, 1)
            table.style = 'Table Grid'
            cell = table.cell(0, 0)
            cell.text = i.removeprefix('```')
            
            # Set background color to dark gray and text color to black
            run = cell.paragraphs[0].runs[0]
            run.font.color.rgb = docx.shared.RGBColor(255, 255, 255)
            run.bold = True
            shading_elm = docx.oxml.parse_xml(r'<w:shd {} w:fill="0F172A"/>'.format(docx.oxml.ns.nsdecls('w')))
            cell._element.get_or_add_tcPr().append(shading_elm)
        else:
            doc.add_paragraph((i).replace("*","•"))
            
    if  doc.save('output.docx'):
        print("Document saved successfully")
    return True
    
# GUI Elements
heading = Label(app, text="Docxer", font=("Arial", 16))
fileName = Label(app, textvariable=inputState, state="disabled", border=1, relief="solid", width=30, pady=6, background=fieldColor, fg="white")
fileName.configure(bg=fieldColor, highlightbackground=fieldColor, highlightcolor=fieldColor, highlightthickness=1)
selectBtn = Button(app, text="Select File", command= lambda: selectFile())
logsLabel = Label(app, textvariable=logs)

# Placing GUI Elements
heading.place(x=90, y=20)
fileName.place(x=20, y=70)
selectBtn.place(x=95, y=130)
logsLabel.place(x=40, y=225)

# Run the app
if __name__ == '__main__':
    app.mainloop()